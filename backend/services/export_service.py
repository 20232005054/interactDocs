"""
文档导出服务

支持格式：Word (.docx)、PDF、Markdown

段落内容存储格式为纯 Markdown 字符串。
- md_to_html：Markdown → HTML，供 PDF（weasyprint）使用
- md_to_text：Markdown → 纯文本，供 Word（python-docx）使用
- extract_image_urls_from_md：从 Markdown ![](url) 语法提取图片 URL
"""

import io
import re
from typing import Optional
from uuid import UUID

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select

from db.models import Chapter, Paragraph, DocumentSummary, Document


# ---------------------------------------------------------------------------
# 内容处理工具函数（Markdown 为输入格式）
# ---------------------------------------------------------------------------

def md_to_html(content: str) -> str:
    """Markdown → HTML，供 PDF 渲染使用"""
    if not content:
        return ""
    import markdown as md_lib
    return md_lib.markdown(
        content,
        extensions=["tables", "fenced_code", "nl2br"],
    )


def md_to_text(content: str) -> str:
    """Markdown → 纯文本，去除所有标记符号，供 Word 段落文本使用"""
    if not content:
        return ""
    # 先转 HTML，再用 BeautifulSoup 提取纯文本，保留换行结构
    from bs4 import BeautifulSoup
    html = md_to_html(content)
    return BeautifulSoup(html, "html.parser").get_text("\n").strip()


def extract_image_urls_from_md(content: str) -> list[str]:
    """从 Markdown ![alt](url) 语法中提取图片 URL"""
    if not content:
        return []
    return re.findall(r'!\[[^\]]*\]\(([^)]+)\)', content)


# ---------------------------------------------------------------------------
# 文档数据组装（格式无关）
# ---------------------------------------------------------------------------

async def _build_references(db: AsyncSession, document_id: UUID) -> list[dict]:
    """获取文档参考文献列表，格式化为温哥华引文"""
    try:
        from services.langchain.services.literature_rag_service import LiteratureRagService
        citations = await LiteratureRagService.get_document_reference_list(db, document_id)
        return [
            {
                "number": c["citation_number"],
                "formatted": LiteratureRagService.format_vancouver_reference(c, c["citation_number"]),
            }
            for c in citations
        ]
    except Exception:
        return []


async def build_document_data(db: AsyncSession, document_id: UUID) -> dict:
    """组装文档完整数据，返回格式无关的中间结构"""
    # 文档基本信息
    doc_result = await db.execute(select(Document).where(Document.document_id == document_id))
    document = doc_result.scalar_one_or_none()
    if not document:
        return None

    # 摘要
    summaries_result = await db.execute(
        select(DocumentSummary)
        .where(DocumentSummary.document_id == document_id)
        .order_by(DocumentSummary.order_index)
    )
    summaries = summaries_result.scalars().all()

    # 章节（全部加载，在内存中构建树）
    chapters_result = await db.execute(
        select(Chapter)
        .where(Chapter.document_id == document_id)
        .order_by(Chapter.order_index)
    )
    chapters = chapters_result.scalars().all()

    # 段落（批量加载）
    chapter_ids = [c.chapter_id for c in chapters]
    paragraphs_result = await db.execute(
        select(Paragraph)
        .where(Paragraph.chapter_id.in_(chapter_ids))
        .order_by(Paragraph.chapter_id, Paragraph.order_index)
    )
    paragraphs = paragraphs_result.scalars().all()

    # 按 chapter_id 分组段落
    para_map: dict[UUID, list] = {}
    for p in paragraphs:
        para_map.setdefault(p.chapter_id, []).append(p)

    # 构建章节树
    chapter_map = {c.chapter_id: c for c in chapters}

    def build_chapter_node(chapter: Chapter, level: int) -> dict:
        return {
            "chapter_id": str(chapter.chapter_id),
            "parent_id": str(chapter.parent_id) if chapter.parent_id else None,
            "title": chapter.title,
            "order_index": chapter.order_index,
            "level": level,
            "paragraphs": [
                {
                    "content": p.content,
                    "para_type": p.para_type,
                    "order_index": p.order_index,
                }
                for p in para_map.get(chapter.chapter_id, [])
            ],
            "children": [],
        }

    # 先建所有节点
    nodes = {c.chapter_id: build_chapter_node(c, 1) for c in chapters}

    # 计算层级并构建树
    root_chapters = []
    for chapter in chapters:
        node = nodes[chapter.chapter_id]
        if chapter.parent_id and chapter.parent_id in nodes:
            parent_node = nodes[chapter.parent_id]
            node["level"] = parent_node["level"] + 1
            parent_node["children"].append(node)
        else:
            root_chapters.append(node)

    return {
        "title": document.title,
        "purpose": document.purpose or "",
        "summaries": [
            {"title": s.title, "content": s.content}
            for s in summaries
        ],
        "chapters": root_chapters,
        "references": await _build_references(db, document_id),
    }


# ---------------------------------------------------------------------------
# Word 导出
# ---------------------------------------------------------------------------

def _set_run_font(run, font_name: str = "宋体"):
    """为 run 设置中文字体，避免中文显示为方块"""
    from docx.oxml.ns import qn
    from lxml import etree
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_name)


def _render_inline(paragraph, node):
    """递归渲染行内节点（加粗、斜体、普通文本）到 docx paragraph"""
    from bs4 import NavigableString, Tag
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            _set_run_font(paragraph.add_run(text))
        return
    if not isinstance(node, Tag):
        return
    tag = node.name
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = paragraph.add_run(text)
                if tag in ("strong", "b"):
                    run.bold = True
                elif tag in ("em", "i"):
                    run.italic = True
                _set_run_font(run)
        else:
            _render_inline(paragraph, child)


def _add_md_table(doc, table_node):
    """将 BeautifulSoup table 节点写入 docx Table"""
    from docx.shared import Pt
    rows = table_node.find_all("tr")
    if not rows:
        return
    # 计算列数
    col_count = max(len(r.find_all(["td", "th"])) for r in rows)
    if col_count == 0:
        return
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for c_idx, cell in enumerate(cells):
            if c_idx >= col_count:
                break
            docx_cell = tbl.cell(r_idx, c_idx)
            text = cell.get_text(strip=True)
            p = docx_cell.paragraphs[0]
            run = p.add_run(text)
            if cell.name == "th":
                run.bold = True
            _set_run_font(run)


async def _render_md_to_docx(doc, content: str):
    """
    将 Markdown 内容渲染到 docx Document。
    流程：Markdown → HTML → BeautifulSoup → python-docx 对象
    图片异步下载（asyncio.to_thread 包装）。
    """
    import asyncio
    from bs4 import BeautifulSoup, NavigableString, Tag
    from docx.shared import Inches

    if not content or not content.strip():
        return

    html = md_to_html(content)
    soup = BeautifulSoup(html, "html.parser")

    for node in soup.children:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                p = doc.add_paragraph()
                _set_run_font(p.add_run(text))
            continue
        if not isinstance(node, Tag):
            continue

        tag = node.name

        # 标题
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            h = doc.add_heading(node.get_text(strip=True), level=level)
            for run in h.runs:
                _set_run_font(run)

        # 段落
        elif tag == "p":
            # 检查是否含图片
            imgs = node.find_all("img")
            if imgs:
                for img in imgs:
                    url = img.get("src", "")
                    if url:
                        try:
                            import httpx
                            resp = await asyncio.to_thread(
                                lambda u=url: httpx.get(u, timeout=10, follow_redirects=True)
                            )
                            if resp.status_code == 200:
                                doc.add_picture(io.BytesIO(resp.content), width=Inches(5))
                        except Exception:
                            pass
                # 图片后的文字
                text = node.get_text(strip=True)
                if text:
                    p = doc.add_paragraph()
                    _render_inline(p, node)
            else:
                p = doc.add_paragraph()
                _render_inline(p, node)

        # 表格
        elif tag == "table":
            _add_md_table(doc, node)

        # 无序列表
        elif tag == "ul":
            for li in node.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet")
                _render_inline(p, li)

        # 有序列表
        elif tag == "ol":
            for li in node.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Number")
                _render_inline(p, li)

        # 代码块
        elif tag == "pre":
            code = node.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = "Courier New"

        # 水平线
        elif tag == "hr":
            doc.add_paragraph("─" * 40)


async def export_docx(db: AsyncSession, document_id: UUID) -> bytes:
    """生成 Word 文档，返回二进制内容"""
    from docx import Document as DocxDocument

    data = await build_document_data(db, document_id)
    if not data:
        raise ValueError("文档不存在")

    doc = DocxDocument()
    title_heading = doc.add_heading(data["title"], level=0)
    for run in title_heading.runs:
        _set_run_font(run)

    if data["purpose"]:
        p = doc.add_paragraph()
        _set_run_font(p.add_run(f"用途：{data['purpose']}"))

    # 摘要
    if data["summaries"]:
        summary_heading = doc.add_heading("摘要信息", level=1)
        for run in summary_heading.runs:
            _set_run_font(run)
        for s in data["summaries"]:
            title_p = doc.add_paragraph()
            run_title = title_p.add_run(f"{s['title']}")
            run_title.bold = True
            _set_run_font(run_title)
            await _render_md_to_docx(doc, s["content"] or "")

    # 章节递归写入
    async def write_chapter(chapter: dict):
        level = min(chapter["level"], 9)
        heading = doc.add_heading(chapter["title"], level=level)
        for run in heading.runs:
            _set_run_font(run)

        for para in chapter["paragraphs"]:
            content = para["content"] or ""
            para_type = para.get("para_type", "paragraph")

            if para_type in ("heading1", "heading2", "heading3"):
                h_level = int(para_type[-1])
                # 内容可能已含 # 前缀，提取纯文本
                text = content.lstrip("#").strip()
                h = doc.add_heading(text, level=h_level)
                for run in h.runs:
                    _set_run_font(run)
            else:
                await _render_md_to_docx(doc, content)

        for child in chapter.get("children", []):
            await write_chapter(child)

    for chapter in data["chapters"]:
        await write_chapter(chapter)

    # 参考文献
    if data.get("references"):
        ref_heading = doc.add_heading("参考文献", level=1)
        for run in ref_heading.runs:
            _set_run_font(run)
        for ref in data["references"]:
            p = doc.add_paragraph()
            _set_run_font(p.add_run(ref["formatted"]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF 导出
# ---------------------------------------------------------------------------

async def export_pdf(db: AsyncSession, document_id: UUID) -> bytes:
    """生成 PDF，返回二进制内容（weasyprint 渲染 HTML）"""
    import weasyprint

    data = await build_document_data(db, document_id)
    if not data:
        raise ValueError("文档不存在")

    def chapter_to_html(chapter: dict) -> str:
        level = min(chapter["level"] + 1, 6)
        html = f"<h{level}>{chapter['title']}</h{level}>\n"
        for para in chapter["paragraphs"]:
            content = para["content"] or ""
            para_type = para.get("para_type", "paragraph")
            if para_type in ("heading1", "heading2", "heading3"):
                h_level = int(para_type[-1])
                # 内容本身是 Markdown，转 HTML 后直接输出
                html += md_to_html(content) + "\n"
            else:
                # 正文段落：Markdown → HTML，直接输出（不套额外 <p>）
                html += md_to_html(content) + "\n"
        for child in chapter.get("children", []):
            html += chapter_to_html(child)
        return html

    # 摘要：内容是 Markdown，转 HTML 后放入表格
    summaries_html = ""
    if data["summaries"]:
        summaries_html = "<h2>摘要信息</h2>\n<table>\n"
        for s in data["summaries"]:
            content_html = md_to_html(s["content"] or "")
            summaries_html += f"<tr><td class='summary-title'><b>{s['title']}</b></td><td>{content_html}</td></tr>\n"
        summaries_html += "</table>\n"

    chapters_html = "".join(chapter_to_html(c) for c in data["chapters"])

    references_html = ""
    if data.get("references"):
        references_html = "<h2>参考文献</h2>\n<ol>\n"
        for ref in data["references"]:
            references_html += f"<li>{ref['formatted']}</li>\n"
        references_html += "</ol>\n"

    purpose_html = f"<p class='purpose'>用途：{data['purpose']}</p>" if data["purpose"] else ""

    full_html = f"""<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<style>
  body {{ font-family: 'SimSun', serif; font-size: 12pt; margin: 2cm; line-height: 1.8; }}
  h1 {{ font-size: 20pt; text-align: center; margin-bottom: 0.5em; }}
  h2 {{ font-size: 15pt; margin-top: 1.5em; border-bottom: 1px solid #ccc; padding-bottom: 4px; }}
  h3 {{ font-size: 13pt; margin-top: 1em; }}
  h4, h5, h6 {{ font-size: 12pt; margin-top: 0.8em; }}
  p {{ margin: 0.4em 0; text-indent: 2em; }}
  table {{ border-collapse: collapse; width: 100%; margin: 0.8em 0; }}
  td, th {{ border: 1px solid #999; padding: 6px 8px; vertical-align: top; }}
  th {{ background: #f0f0f0; font-weight: bold; }}
  .summary-title {{ width: 8em; white-space: nowrap; }}
  .purpose {{ color: #666; font-style: italic; text-indent: 0; }}
  ol, ul {{ margin: 0.4em 0; padding-left: 2em; }}
  li {{ margin: 0.2em 0; }}
  code {{ font-family: monospace; background: #f5f5f5; padding: 1px 4px; }}
  pre {{ background: #f5f5f5; padding: 8px; overflow-x: auto; }}
</style>
</head><body>
<h1>{data['title']}</h1>
{purpose_html}
{summaries_html}
{chapters_html}
{references_html}
</body></html>"""

    pdf_bytes = weasyprint.HTML(string=full_html).write_pdf()
    return pdf_bytes


# ---------------------------------------------------------------------------
# Markdown 导出
# ---------------------------------------------------------------------------

async def export_markdown(db: AsyncSession, document_id: UUID) -> str:
    """生成 Markdown 文本"""
    data = await build_document_data(db, document_id)
    if not data:
        raise ValueError("文档不存在")

    lines = [f"# {data['title']}\n"]

    if data["purpose"]:
        lines.append(f"> 用途：{data['purpose']}\n")

    # 摘要：内容本身已是 Markdown，直接输出
    if data["summaries"]:
        lines.append("## 摘要信息\n")
        for s in data["summaries"]:
            content = (s["content"] or "").strip()
            lines.append(f"**{s['title']}**\n\n{content}\n")

    def chapter_to_md(chapter: dict) -> list[str]:
        prefix = "#" * min(chapter["level"] + 1, 6)
        result = [f"{prefix} {chapter['title']}\n"]
        for para in chapter["paragraphs"]:
            content = (para["content"] or "").strip()
            para_type = para.get("para_type", "paragraph")
            if para_type in ("heading1", "heading2", "heading3"):
                # 段落内嵌标题：直接输出（内容本身可能已含 # 前缀，也可能是纯文本）
                h_level = int(para_type[-1])
                # 若内容已有 # 前缀则直接用，否则补上
                if content.startswith("#"):
                    result.append(f"{content}\n")
                else:
                    result.append(f"{'#' * h_level} {content}\n")
            else:
                # 正文段落：内容本身是 Markdown，直接输出
                if content:
                    result.append(f"{content}\n")
        for child in chapter.get("children", []):
            result.extend(chapter_to_md(child))
        return result

    for chapter in data["chapters"]:
        lines.extend(chapter_to_md(chapter))

    # 参考文献
    if data.get("references"):
        lines.append("\n## 参考文献\n")
        for ref in data["references"]:
            lines.append(f"{ref['formatted']}\n")

    return "\n".join(lines)
